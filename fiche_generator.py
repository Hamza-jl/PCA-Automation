"""
fiche_generator.py — Generates one BIA fiche per department/structure
from a 'fiche de recensement des structures' (Suivi projet.xlsx) and a
BIA fiche template (.docx).

Entry point:
    generate_all_fiches(
        xlsx_path, template_path, output_dir,
        version="2.0", openai_api_key=None
    ) -> (list[Path], list[str])
"""

from __future__ import annotations

import base64
import json
import os
import re
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Optional

import openpyxl
from docx import Document


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _clean(value) -> str:
    """Normalise a cell value to a stripped single-line string."""
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value).replace("\n", " ")).strip()


def _safe_filename(name: str) -> str:
    """Strip characters illegal in Windows file names."""
    return re.sub(r'[\\/:*?"<>|]', "-", name).strip(" .")


# ---------------------------------------------------------------------------
# Extract ALL images from the recensement Excel (via zip — no BytesIO bugs)
# ---------------------------------------------------------------------------

def _extract_all_images(xlsx_path: Path) -> list[tuple[str, bytes]]:
    """
    Returns [(filename, raw_bytes), ...] for every image in xl/media/.
    Reading from the zip directly avoids the openpyxl BytesIO seek bug.
    """
    try:
        with zipfile.ZipFile(str(xlsx_path), "r") as z:
            media = [
                f for f in z.namelist()
                if re.match(r"xl/media/.*\.(png|jpg|jpeg|gif|bmp)$", f, re.IGNORECASE)
            ]
            return [(f, z.read(f)) for f in sorted(media)]
    except Exception:
        return []


# ---------------------------------------------------------------------------
# OpenAI Vision — detect client name + identify the client logo
# ---------------------------------------------------------------------------

def detect_client_info(
    xlsx_path: Path,
    api_key: str,
) -> tuple[str, Optional[bytes]]:
    """
    Uses GPT-4o Vision to:
      1. Identify which image in the xlsx is the CLIENT logo (not Devoteam).
      2. Extract the client company name from that logo.

    Returns (client_name, logo_bytes).
    Falls back to ("Client", largest_image) if the API call fails.
    """
    images = _extract_all_images(xlsx_path)
    if not images:
        return "Client", None

    # Fallback: largest image (used if API call fails)
    fallback_logo = max(images, key=lambda x: len(x[1]))[1]

    try:
        import openai

        client = openai.OpenAI(api_key=api_key)

        # Build the message: text instruction + one image_url block per image
        content: list[dict] = [
            {
                "type": "text",
                "text": (
                    "The following images are extracted from an Excel file used for a BIA "
                    "(Business Impact Analysis) project managed by Devoteam (a consulting firm).\n\n"
                    "Please analyse every image and:\n"
                    "1. Identify the CLIENT company logo — this is NOT Devoteam's logo.\n"
                    "2. Extract the exact client company name as it appears in the logo.\n\n"
                    "Respond ONLY with a valid JSON object, no markdown, no explanation:\n"
                    '{"client_name": "<name>", "logo_index": <0-based index or null>}'
                ),
            }
        ]

        for fname, data in images:
            ext = fname.rsplit(".", 1)[-1].lower()
            mime = "image/jpeg" if ext in ("jpg", "jpeg") else f"image/{ext}"
            b64 = base64.b64encode(data).decode()
            content.append(
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{mime};base64,{b64}",
                        "detail": "low",
                    },
                }
            )

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": content}],
            max_tokens=150,
        )

        raw = response.choices[0].message.content.strip()
        # Strip markdown code fences if present
        raw = re.sub(r"^```[a-z]*\n?", "", raw).rstrip("` \n")
        result = json.loads(raw)

        client_name = str(result.get("client_name") or "Client").strip()
        logo_index  = result.get("logo_index")

        logo_bytes: Optional[bytes] = None
        if logo_index is not None and 0 <= int(logo_index) < len(images):
            logo_bytes = images[int(logo_index)][1]
        else:
            logo_bytes = fallback_logo

        return client_name, logo_bytes

    except Exception as exc:
        # Non-fatal — return a safe default so generation still runs
        print(f"[fiche_generator] OpenAI detection failed: {exc}")
        return "Client", fallback_logo


# ---------------------------------------------------------------------------
# Parse the recensement sheet into a list of structures
# ---------------------------------------------------------------------------

def parse_structures(xlsx_path: Path) -> list[dict]:
    """
    Reads the 'Planning des réunions' sheet and returns one dict per row
    that represents a concrete structure (department / sub-department).

    Each dict:
        name       – most specific level name (N3 > N2 > N1)
        vis_a_vis  – contact person name
        date       – meeting date as "DD/MM/YYYY" string
        niveau1    – parent Pôle (cascaded)
        niveau2    – parent Département (cascaded)
        niveau3    – sub-department (may be empty)
    """
    wb = openpyxl.load_workbook(str(xlsx_path), data_only=True)
    sheet_name = next(
        (s for s in wb.sheetnames
         if "réunion" in s.lower() or "reunion" in s.lower()),
        wb.sheetnames[0],
    )
    ws = wb[sheet_name]

    COL_N1, COL_N2, COL_N3 = 0, 1, 2
    COL_VAV, COL_DATE       = 3, 4

    header_found = False
    structures: list[dict] = []
    current_n1 = ""
    current_n2 = ""

    for row in ws.iter_rows(values_only=True):
        vals = list(row)

        if not header_found:
            if vals[0] and "niveau" in str(vals[0]).lower():
                header_found = True
            continue

        n1  = _clean(vals[COL_N1]) if COL_N1 < len(vals) else ""
        n2  = _clean(vals[COL_N2]) if COL_N2 < len(vals) else ""
        n3  = _clean(vals[COL_N3]) if COL_N3 < len(vals) else ""
        vav = _clean(vals[COL_VAV]) if COL_VAV < len(vals) else ""
        date_raw = vals[COL_DATE]   if COL_DATE < len(vals) else None

        if n1:
            current_n1 = n1
        if n2:
            current_n2 = n2

        name = n3 or n2 or n1
        if not name:
            continue

        date_str = ""
        if date_raw:
            if isinstance(date_raw, datetime):
                date_str = date_raw.strftime("%d/%m/%Y")
            else:
                date_str = str(date_raw)

        structures.append(
            {
                "name":    name,
                "vis_a_vis": vav,
                "date":    date_str,
                "niveau1": current_n1,
                "niveau2": current_n2,
                "niveau3": n3,
            }
        )

    return structures


# ---------------------------------------------------------------------------
# Text replacement helpers
# ---------------------------------------------------------------------------

def _replace_in_para(para, old: str, new: str) -> None:
    """
    Replace *old* with *new* in *para* while preserving run formatting.
    Tries run-by-run first; falls back to collapsing all text into run[0].
    """
    if old not in para.text:
        return
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return
    # Fallback: merge all text into first run
    new_full = para.text.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for run in para.runs[1:]:
            run.text = ""


def _replace_in_cell(cell, old: str, new: str) -> None:
    for para in cell.paragraphs:
        _replace_in_para(para, old, new)


def _replace_in_doc(doc: Document, old: str, new: str) -> None:
    """Replace *old* with *new* in every paragraph of the document body."""
    for para in doc.paragraphs:
        _replace_in_para(para, old, new)


# ---------------------------------------------------------------------------
# Logo replacement in header
# ---------------------------------------------------------------------------

def _replace_header_logo(doc: Document, logo_bytes: bytes) -> None:
    """
    Replaces the client logo (rId1) in each non-linked header.
    The Devoteam logo (rId2) is left untouched.

    Confirmed mapping for the STAR template:
      rId1 → image1.png  = client / STAR logo   ← replace
      rId2 → image2.png  = Devoteam logo         ← keep
    """
    for section in doc.sections:
        if section.header.is_linked_to_previous:
            continue
        part = section.header.part
        if "rId1" not in part.rels:
            continue
        # Overwrite blob in-place (content_type is read-only on ImagePart)
        part.rels["rId1"]._target._blob = logo_bytes


# ---------------------------------------------------------------------------
# Generate one fiche
# ---------------------------------------------------------------------------

def generate_fiche(
    template_path: Path,
    structure: dict,
    logo_bytes: Optional[bytes],
    output_path: Path,
    client_name: str = "Client",
    version: str = "2.0",
) -> None:
    """
    Clone the BIA template, fill entity-specific fields, replace the client
    name and logo, and save to *output_path*.
    Blank sections are intentionally left blank.
    """
    doc = Document(str(template_path))

    name     = structure["name"]
    vis_a_vis = structure["vis_a_vis"]
    date_str = structure["date"]
    doc_ref  = f"{client_name} - MCO - BIA - {name} - V{version}"

    # ------------------------------------------------------------------
    # 1. Replace hard-coded client name ("STAR Assurances") everywhere
    # ------------------------------------------------------------------
    _replace_in_doc(doc, "STAR Assurances", client_name)
    _replace_in_doc(doc, "STAR",            client_name)  # catch remaining

    # ------------------------------------------------------------------
    # 2. Cover page paragraph: "Entité : xx"
    # ------------------------------------------------------------------
    for para in doc.paragraphs:
        if "Entité" in para.text and "xx" in para.text:
            _replace_in_para(para, "xx", name)
            break

    # ------------------------------------------------------------------
    # 3. Identity table (Table index 2)
    # ------------------------------------------------------------------
    try:
        t = doc.tables[2]

        # Row 0 — entity name
        _replace_in_cell(t.cell(0, 1), "xx", name)
        _replace_in_cell(t.cell(0, 2), "xx", name)

        # Row 9 — version
        _replace_in_cell(t.cell(9, 1), "2.0", version)
        _replace_in_cell(t.cell(9, 2), "2.0", version)

        # Row 10 — date of last update
        if date_str:
            _replace_in_cell(t.cell(10, 1), "xx", date_str)
            _replace_in_cell(t.cell(10, 2), "xx", date_str)

        # Row 11 — document reference
        _replace_in_cell(t.cell(11, 1), "xx", doc_ref)
        _replace_in_cell(t.cell(11, 2), "xx", doc_ref)

        # Row 2 — first attendee (the client vis-à-vis)
        valid_vav = vis_a_vis and vis_a_vis.lower() not in ("à définir", "a definir", "")
        if valid_vav:
            _replace_in_cell(t.cell(2, 1), "Mr Lazher HEDFI", vis_a_vis)

    except Exception:
        pass

    # ------------------------------------------------------------------
    # 4. "Tableau 6 : xx" caption
    # ------------------------------------------------------------------
    for para in doc.paragraphs:
        if "Tableau 6" in para.text and "xx" in para.text:
            _replace_in_para(para, "xx", name)

    # ------------------------------------------------------------------
    # 5. Replace client logo in every non-linked header
    # ------------------------------------------------------------------
    if logo_bytes:
        _replace_header_logo(doc, logo_bytes)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def generate_all_fiches(
    xlsx_path: Path,
    template_path: Path,
    output_dir: Path,
    version: str = "2.0",
    openai_api_key: Optional[str] = None,
) -> tuple[list[Path], list[str]]:
    """
    Generate one BIA fiche per structure found in *xlsx_path*.

    Args:
        xlsx_path       – recensement Excel file
        template_path   – blank BIA fiche template (.docx)
        output_dir      – folder to write generated fiches into
        version         – "1.0" or "2.0"
        openai_api_key  – GPT-4o key for client detection; falls back to
                          the OPENAI_API_KEY env var if not provided

    Returns:
        (generated_paths, error_strings)
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    # Resolve API key
    api_key = openai_api_key or os.environ.get("OPENAI_API_KEY", "")

    # Use OpenAI to identify client name + logo; fall back gracefully
    if api_key:
        client_name, logo_bytes = detect_client_info(xlsx_path, api_key)
    else:
        # No key — use largest image, unknown client name
        images = _extract_all_images(xlsx_path)
        logo_bytes = max(images, key=lambda x: len(x[1]))[1] if images else None
        client_name = "Client"

    structures = parse_structures(xlsx_path)

    generated: list[Path] = []
    errors: list[str] = []

    for s in structures:
        safe_name = _safe_filename(s["name"])
        safe_client = _safe_filename(client_name)
        out_path = output_dir / f"{safe_client} - MCO - BIA - {safe_name} - V{version}.docx"
        try:
            generate_fiche(
                template_path, s, logo_bytes, out_path,
                client_name=client_name,
                version=version,
            )
            generated.append(out_path)
        except Exception as exc:
            errors.append(f"{s['name']}: {exc}")

    return generated, errors
